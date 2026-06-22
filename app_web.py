# -*- coding: utf-8 -*-
"""SiC Lab · Flask Web App with Auth"""
import os, sys, time, tempfile, re, json
from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from orchestrator import run_orchestrator_pipeline
from functools import wraps
from concurrent.futures import ThreadPoolExecutor

app = Flask(__name__)
app.secret_key = os.environ.get('SIC_SECRET', 'sic-lab-' + str(int(time.time())))
ACCESS_KEY = os.environ.get('SIC_LAB_KEY', 'siclab2026')

# Auto-load DeepSeek API key from OpenCode config
def _load_deepseek_key():
    try:
        auth_path = os.path.join(os.path.expanduser('~'), '.local', 'share', 'opencode', 'auth.json')
        if os.path.exists(auth_path):
            with open(auth_path, 'r', encoding='utf-8') as f:
                auth = json.load(f)
            return auth.get('deepseek', {}).get('key', '')
    except: pass
    return os.environ.get('DEEPSEEK_API_KEY', '')

DEEPSEEK_KEY = _load_deepseek_key()

UPLOAD_DIR = tempfile.mkdtemp(prefix='sic_')

# Paper cache: {filepath_mtime: extracted_data}
paper_cache = {}

# ============ Chat Database ============
import sqlite3
CHAT_DB = os.path.join(os.path.expanduser('~'), '.siclab_chat.db')

def init_chat_db():
    with sqlite3.connect(CHAT_DB) as db:
        db.execute('''CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT DEFAULT 'New Chat',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
        db.execute('''CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conv_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (conv_id) REFERENCES conversations(id)
        )''')
        db.commit()

init_chat_db()

def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('authenticated'):
            if request.path.startswith('/api/'):
                return jsonify({'error': 'unauthorized'}), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        if request.form.get('key', '') == ACCESS_KEY:
            session['authenticated'] = True
            return redirect(url_for('index'))
        error = '密钥错误，请重试'
    return render_template('login.html', error=error)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/')
@require_auth
def index():
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
@require_auth
def upload():
    files = request.files.getlist('papers')
    results = []
    for f in files:
        if f.filename.lower().endswith('.pdf'):
            name = secure_filename(f.filename)
            path = os.path.join(UPLOAD_DIR, name)
            f.save(path)
            results.append({'name': f.filename, 'path': path, 'size': os.path.getsize(path)})
    return jsonify({'count': len(results), 'files': results})

@app.route('/api/analyze', methods=['POST'])
@require_auth
def analyze():
    data = request.get_json()
    uploads = data.get('uploads', [])
    goal = data.get('goal', '').strip()
    sample_status = data.get('sampleStatus', '').strip()
    requirements = data.get('requirements', '').strip()
    all_text = data.get('allText', '').strip()
    
    # Run Paper-Orchestrator pipeline
    report, papers = run_orchestrator_pipeline(uploads, sample_status, goal, requirements, DEEPSEEK_KEY)
    
    # Generate Word document
    doc_path = generate_word_doc(papers, report, all_text)
    
    return jsonify({'plan': report, 'download': os.path.basename(doc_path)})

@app.route('/api/refine', methods=['POST'])
@require_auth
def refine_report():
    data = request.get_json()
    report = data.get('report', '')
    feedback = data.get('feedback', '')
    papers = data.get('papers', [])
    
    if not DEEPSEEK_KEY:
        return jsonify({'report': report, 'reply': '无法连接AI，请检查API密钥。'})
    
    try:
        from openai import OpenAI
        client = OpenAI(api_key=DEEPSEEK_KEY, base_url='https://api.deepseek.com/v1')
        resp = client.chat.completions.create(
            model='deepseek-v4-pro',
            messages=[{'role': 'system', 'content': '''You are a STRICT report editor. Follow user instructions EXACTLY.

RULES:
1. If user says "delete all #" - REMOVE EVERY # character from the entire report. No exceptions.
2. If user says "remove XXX" - ensure XXX does NOT appear in the final report.
3. If user says "change A to B" - change ALL occurrences.
4. Be thorough - check the entire report for compliance.
5. Explain what you changed in 1-2 Chinese sentences.
6. If the request is unclear, ask a specific clarifying question.
7. After explaining, output the complete fixed report after "---REPORT---".

Format:
[Brief explanation in Chinese]
---REPORT---
[Complete modified report with all changes applied]'''},
                      {'role': 'user', 'content': f'Current report:\n{report}\n\nUser feedback: {feedback}'}],
            temperature=0.4, max_tokens=4000
        )
        content = resp.choices[0].message.content
        
        # Split reply and report
        if '---REPORT---' in content:
            parts = content.split('---REPORT---', 1)
            reply = parts[0].strip()
            new_report = parts[1].strip()
        else:
            reply = '已根据你的要求修改报告。'
            new_report = content
        
        # Generate new Word doc
        doc_path = ''
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            doc = Document()
            style = doc.styles['Normal']; style.font.size = Pt(11)
            style.font.name = 'Times New Roman'
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = style.element.rPr.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), 'SimSun')
            
            for _ in range(4): doc.add_paragraph('')
            t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = t.add_run('Advanced Materials\n实验报告'); r.font.size = Pt(26); r.font.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            doc.add_page_break()
            for line in new_report.split('\n'):
                line = line.strip()
                if not line: continue
                doc.add_paragraph(line)
            
            doc_path = os.path.join(tempfile.gettempdir(), f'SiC_Refined_{int(time.time())}.docx')
            doc.save(doc_path)
            doc_path = os.path.basename(doc_path)
        except:
            pass
        
        return jsonify({'report': new_report, 'reply': reply, 'download': doc_path})
    except:
        return jsonify({'report': report, 'reply': '修改失败，请重试。'})

@app.route('/api/export-pdf', methods=['POST'])
@require_auth
def export_pdf():
    data = request.get_json()
    content = data.get('content', '')
    
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('SimSun', '', r'C:\Windows\Fonts\simsun.ttc', uni=True)
    pdf.set_font('SimSun', '', 11)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        pdf.multi_cell(0, 6, line)
    
    path = os.path.join(tempfile.gettempdir(), f'AM_Report_{int(time.time())}.pdf')
    pdf.output(path)
    return jsonify({'download': os.path.basename(path)})

@app.route('/api/analyze-image', methods=['POST'])
@require_auth
def analyze_image():
    if 'image' not in request.files:
        return jsonify({'result': 'No image uploaded'})
    
    file = request.files['image']
    img_path = os.path.join(tempfile.gettempdir(), 'img_' + secure_filename(file.filename))
    file.save(img_path)
    
    # Try DeepSeek vision
    if DEEPSEEK_KEY:
        try:
            import base64
            with open(img_path, 'rb') as f:
                img_b64 = base64.b64encode(f.read()).decode()
            
            from openai import OpenAI
            client = OpenAI(api_key=DEEPSEEK_KEY, base_url='https://api.deepseek.com/v1')
            resp = client.chat.completions.create(
                model='deepseek-v4-pro',
                messages=[{'role': 'user', 'content': [
                    {'type': 'text', 'text': 'Analyze this materials science image (SEM/XRD/metallography). Describe what you see: morphology, phases, grain size, defects, features. Be specific and quantitative. Use Chinese if the content appears to be Chinese research.'},
                    {'type': 'image_url', 'image_url': {'url': f'data:image/jpeg;base64,{img_b64}'}}
                ]}],
                max_tokens=800
            )
            return jsonify({'result': resp.choices[0].message.content})
        except Exception as e:
            return jsonify({'result': f'AI analysis unavailable: {str(e)[:100]}'})
    
    return jsonify({'result': f'Image received: {file.filename} ({os.path.getsize(img_path)} bytes). AI key required for analysis.'})

# ============ Knowledge Base ============
KB_DIR = os.path.join(os.path.expanduser('~'), 'AdvancedMaterials_KB')
os.makedirs(KB_DIR, exist_ok=True)
KB_DB = os.path.join(KB_DIR, 'kb_index.json')

def load_kb_index():
    if os.path.exists(KB_DB):
        with open(KB_DB, 'r', encoding='utf-8') as f: return json.load(f)
    return []

def save_kb_index(idx):
    with open(KB_DB, 'w', encoding='utf-8') as f: json.dump(idx, f, ensure_ascii=False, indent=2)

@app.route('/api/kb/add', methods=['POST'])
@require_auth
def kb_add():
    files = request.files.getlist('papers')
    note = request.form.get('note', '')
    idx = load_kb_index()
    count = 0
    for f in files:
        if f.filename.lower().endswith('.pdf'):
            name = str(int(time.time())) + '_' + secure_filename(f.filename)
            path = os.path.join(KB_DIR, name)
            f.save(path)
            idx.append({'id': len(idx)+1, 'name': f.filename, 'path': path, 'note': note, 'added': time.strftime('%Y-%m-%d'), 'pages': 0})
            count += 1
    save_kb_index(idx)
    return jsonify({'count': count})

@app.route('/api/kb/list')
@require_auth
def kb_list():
    return jsonify(load_kb_index())

@app.route('/api/kb/del/<int:kb_id>', methods=['DELETE'])
@require_auth
def kb_del(kb_id):
    idx = load_kb_index()
    item = next((i for i in idx if i['id'] == kb_id), None)
    if item:
        if os.path.exists(item['path']): os.remove(item['path'])
        idx = [i for i in idx if i['id'] != kb_id]
        save_kb_index(idx)
    return jsonify({'ok': True})

@app.route('/api/kb/import', methods=['POST'])
@require_auth
def kb_import():
    data = request.get_json()
    folder = data.get('folder', '')
    if not folder or not os.path.isdir(folder):
        return jsonify({'count': 0, 'error': 'Invalid folder'})
    
    idx = load_kb_index()
    count = 0
    for root, dirs, files in os.walk(folder):
        for f in files:
            if f.lower().endswith('.pdf'):
                src = os.path.join(root, f)
                name = str(int(time.time() + count)) + '_' + secure_filename(f)
                dst = os.path.join(KB_DIR, name)
                try:
                    import shutil
                    shutil.copy2(src, dst)
                    idx.append({'id': len(idx)+1, 'name': f, 'path': dst, 'note': '批量导入', 'added': time.strftime('%Y-%m-%d'), 'pages': 0})
                    count += 1
                except: pass
    save_kb_index(idx)
    return jsonify({'count': count})

@app.route('/api/download/<filename>')
@require_auth
def download_report(filename):
    if os.path.exists(path):
        return send_file(path, as_attachment=True, download_name='SiC_lab_report.docx')
    return 'Not found', 404

def detect_methods(text):
    patterns = {
        'HP': ['hot press', '热压'], 'SPS': ['spark plasma', '放电等离子'],
        '无压渗透': ['pressureless infiltrat', '无压渗透', '无压浸渗'],
        'RMI': ['reactive melt infiltrat', 'RMI', '反应熔渗'],
        '氧化退火': ['air anneal', 'oxidation anneal', '氧化退火', '空气退火'],
        '氩气退火': ['Ar anneal', 'argon anneal', '氩气退火'],
        '激光熔覆': ['laser clad', '激光熔覆'],
        'LMD': ['laser metal deposit', 'LMD', '激光沉积'],
        'SLM': ['selective laser melt', 'SLM', '选区激光'],
        '3D打印': ['3D print', 'additive manufact', '3D打印', '增材制造'],
        '搅拌铸造': ['stir cast', '搅拌铸造'],
        '挤压铸造': ['squeeze cast', '挤压铸造'],
        '喷射沉积': ['spray deposit', 'spray form', '喷射沉积', 'Osprey'],
        'CVI': ['chemical vapor infiltrat', 'CVI'],
        'PIP': ['precursor infiltrat pyrolysis', 'PIP'],
        '晶须增强': ['whisker', '晶须', 'nanowire'],
        '纤维增强': ['fiber reinforc', '纤维增强', 'Cf/'],
        '稀土助烧': ['rare earth', '稀土', 'Nd2O3', 'Gd2O3', 'Y2O3', 'La2O3', 'Yb2O3'],
        'AlN添加': ['AlN', 'aluminum nitride', '氮化铝'],
        '两步热压': ['two-step', 'two step', '两步', '二步'],
        '液相烧结': ['liquid phase sinter', '液相烧结'],
    }
    found = []
    tl = text.lower()
    for method, keywords in patterns.items():
        for kw in keywords:
            if kw.lower() in tl:
                found.append(method)
                break
    return list(set(found))

def detect_properties(text):
    props = {}
    for m in re.finditer(r'(\d+[\.\d]*)\s*W[/\xb7]\s*m[\xb7-]\s*K', text):
        val = float(m.group(1))
        if 80 < val < 300: props['TC'] = max(props.get('TC', 0), val)
    for m in re.finditer(r'(\d+[\.\d]*)\s*MPa', text):
        val = float(m.group(1))
        if 200 < val < 1200: props['BS'] = max(props.get('BS', 0), val)
    return props

def clean_text(s):
    """Fast text cleaner using regex"""
    import re
    if not s: return ''
    # Keep: Chinese chars, ASCII printable, common punctuation
    return re.sub(r'[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\u2000-\u206f\x20-\x7e\n\r\t\u00b7]', '', s)

def quality_check(text):
    """Check and report quality issues in generated text"""
    issues = []
    if not text: issues.append('Empty output')
    # Check for common garbled characters
    garbled_markers = ['??', '????', ' ', ' ', ' ', ' ']
    for gm in garbled_markers:
        if gm in text:
            issues.append(f'Found garbled marker: {repr(gm)}')
    # Check line endings consistency
    if '\r\r' in text:
        issues.append('Double CR detected')
    # Ensure proper content
    if len(text) < 100:
        issues.append(f'Output too short ({len(text)} chars)')
    return issues

def parse_requirements(sample_status, goal, requirements):
    """Smart requirements parser agent.
    Uses DeepSeek when available, falls back to local rules.
    Returns structured dict."""
    
    raw_text = f'样品状态：{sample_status}\n实验目标：{goal}\n特殊要求：{requirements}'
    result = {
        'author': '',
        'processes': [],
        'targets': [],
        'equipment': [],
        'notes': [],
        'raw': raw_text
    }
    
    # ---- Try DeepSeek for intelligent parsing ----
    if DEEPSEEK_KEY:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=DEEPSEEK_KEY, base_url='https://api.deepseek.com/v1')
            resp = client.chat.completions.create(
                model='deepseek-v4-pro',
                messages=[{'role': 'system', 'content': '''You are a materials science requirements parser. Extract from the user's Chinese text:

返回JSON格式:
{
  "author": "作者姓名，没有则为空字符串",
  "processes": ["需要的工艺，如：氧化退火、铝合金熔渗、SPS烧结、RMI渗硅、热压烧结"],
  "targets": ["性能目标，如：强度700MPa、热导率180W/mK、密度3.0"],
  "equipment": ["需要的设备，如：马弗炉、管式炉、SPS、真空炉"],
  "notes": ["其他注意事项"]
}

注意：从上下文中理解用户的真实意图，不要照搬原文。
只返回JSON，不要其他文字。'''},
                 {'role': 'user', 'content': raw_text}],
                temperature=0, max_tokens=400
            )
            content = resp.choices[0].message.content.strip()
            if '{' in content:
                parsed = json.loads(content[content.index('{'):content.rindex('}')+1])
                result['author'] = parsed.get('author', '')
                result['processes'] = parsed.get('processes', [])
                result['targets'] = parsed.get('targets', [])
                result['equipment'] = parsed.get('equipment', [])
                result['notes'] = parsed.get('notes', [])
                return result
        except: pass
    
    # ---- Local fallback ----
    import re
    all_text = raw_text.lower()
    
    # Author
    for pat in [r'作者[：:是]\s*(\S+)', r'报告[人者][：:是]\s*(\S+)']:
        m = re.search(pat, raw_text)
        if m:
            result['author'] = m.group(1).strip().rstrip('，。,.')[:20]
            break
    
    # Processes
    process_rules = [
        (['渗铝', '铝浸渗', '铝合金渗透'], '铝合金熔渗'),
        (['渗硅', 'rmi'], 'RMI渗硅'),
        (['退火', '氧化退火', 'anneal'], '氧化退火'),
        (['sps', '放电等离子'], 'SPS烧结'),
        (['热压', 'hot press'], '热压烧结'),
    ]
    for keywords, process in process_rules:
        if any(kw in all_text for kw in keywords):
            result['processes'].append(process)
    
    # Targets
    for m in re.finditer(r'(\d+)\s*mpa', all_text):
        result['targets'].append(f'强度{m.group(1)}MPa')
    for m in re.finditer(r'(\d+)\s*w/(?:m|m\.)\s*k', all_text):
        result['targets'].append(f'热导率{m.group(1)}W/mK')
    
    # Equipment
    eq_rules = [
        (['管式炉', 'n2', '气氛'], '管式气氛炉+N2'),
        (['sps'], 'SPS烧结炉'),
        (['真空'], '真空炉'),
        (['马弗炉', '退火'], '马弗炉'),
    ]
    for keywords, eq in eq_rules:
        if any(kw in all_text for kw in keywords):
            result['equipment'].append(eq)
    
    return result


def generate_smart_report(extracted, sample_status, goal, requirements, all_text):
    total = len(extracted)
    status = sample_status.strip() if sample_status else ''
    goals = goal.strip() if goal else ''
    reqs = requirements.strip() if requirements else ''
    
    paper_data = []
    for i, e in enumerate(extracted[:12], 1):
        ms = e.get('methods', [])[:5] if e.get('methods') else []
        ps = e.get('properties', {})
        if not isinstance(ps, dict): ps = {}
        tc = f'TC={ps["TC"]} W/(m.K)' if ps.get('TC') else ''
        bs = f'BS={ps["BS"]} MPa' if ps.get('BS') else ''
        paper_data.append(f'Paper {i}: {e["name"]} ({e["pages"]}p) | Methods: {"; ".join(ms) if ms else "N/A"} | {tc} {bs}')
    
    if DEEPSEEK_KEY:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=DEEPSEEK_KEY, base_url='https://api.deepseek.com/v1')
            resp = client.chat.completions.create(
                model='deepseek-v4-pro',
                messages=[{'role': 'system', 'content': f'''You are a professional materials science experiment report generator. Your task is to generate a comprehensive experiment report based on provided papers and user requirements.

CRITICAL RULES - FOLLOW ALL:
1. Every single user requirement MUST be reflected in the output. Read them ALL carefully.
2. DO NOT copy-paste requirements verbatim. UNDERSTAND them and APPLY them naturally throughout the report.
3. Different papers produce DIFFERENT reports. Use ACTUAL paper data and findings.
4. Work with ANY material (not just SiC). Analyze whatever papers are provided.
5. If user specifies format (font, spacing, language, title-only, etc.), follow EXACTLY.
6. If user says "only a title, no content", give ONLY a title and nothing else.
7. Be SPECIFIC with numbers, temperatures, times, equipment names from the papers.
8. Do NOT add content the user did not ask for. Do NOT use templates.
9. The report MUST feel tailored to THIS specific user and THESE specific papers.
10. Apply requirements intelligently - if user says "English", write in English. If user says "author: Name", include author.

User requirements:
Sample status: {status if status else "Not specified"}
Goal: {goals if goals else "Not specified"}
Special requirements: {reqs if reqs else "None"}

Papers analyzed ({total} total):
{chr(10).join(paper_data)}'''},
                 {'role': 'user', 'content': f'Generate a complete experiment report. Apply ALL my requirements intelligently. Do NOT copy-paste my requirements into the report. My requirements are:\n\n{reqs}\n\nGoal: {goals}\n\nSample: {status}'}],
                temperature=0.3, max_tokens=4000
            )
            return resp.choices[0].message.content
        except:
            pass
    
    return f'Analyzed {total} papers. Goal: {goals}. Requirements: {reqs}'


def generate_word_doc(papers, plan_text, all_text):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    plan_text = clean_text(plan_text)
    doc = Document()
    
    # Extract author from report text (already parsed by AI in generate_smart_report)
    import re
    author = ''
    am = re.search(r'报告作者[：:]\s*(.+)', plan_text)
    if am: author = am.group(1).strip()[:20]
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        hs.font.bold = True
        hs.font.size = {1: Pt(18), 2: Pt(14), 3: Pt(12)}.get(i, Pt(11))
        hrPr = hs.element.get_or_add_rPr()
        hrFonts = hrPr.find(qn('w:rFonts'))
        if hrFonts is None:
            hrFonts = hs.element.rPr.makeelement(qn('w:rFonts'), {})
            hrPr.insert(0, hrFonts)
        hrFonts.set(qn('w:eastAsia'), 'SimHei')
    
    # Cover
    for _ in range(4): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('SiC陶瓷材料实验报告'); r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    doc.add_paragraph('')
    if author:
        ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap.add_run(f'作者：{author}').font.size = Pt(14)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f'分析{len(papers)}篇论文 | {time.strftime("%Y-%m-%d")}').font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_page_break()
    
    for line in plan_text.split('\n'):
        line = line.strip()
        if not line: continue
        if line in ['需求分析', '实验方案']:
            doc.add_heading(line, 1)
        elif line.startswith('详细操作步骤') or line.startswith('性能测试') or line.startswith('注意事项'):
            doc.add_heading(line, 1)
        elif line.startswith('第') and '步' in line:
            doc.add_heading(line, 2)
        elif line.startswith('  ') and line.strip() and not line.startswith('   '):
            doc.add_paragraph(line.strip())
        elif line.startswith('  -'):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    doc.add_paragraph('')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('-- SiC Lab AI platform --').font.color.rgb = RGBColor(0x99,0x99,0x99)
    
    path = os.path.join(tempfile.gettempdir(), f'SiC_Report_{int(time.time())}.docx')
    doc.save(path)
    return path



@app.route('/api/chat/list')
@require_auth
def chat_list():
    with sqlite3.connect(CHAT_DB) as db:
        rows = db.execute('SELECT id, title, created_at FROM conversations ORDER BY created_at DESC').fetchall()
        return jsonify([{'id': r[0], 'title': r[1], 'created_at': r[2]} for r in rows])

@app.route('/api/chat/new', methods=['POST'])
@require_auth
def chat_new():
    with sqlite3.connect(CHAT_DB) as db:
        cur = db.execute('INSERT INTO conversations DEFAULT VALUES')
        return jsonify({'id': cur.lastrowid})

@app.route('/api/chat/delete/<int:conv_id>', methods=['DELETE'])
@require_auth
def chat_delete(conv_id):
    with sqlite3.connect(CHAT_DB) as db:
        db.execute('DELETE FROM messages WHERE conv_id=?', (conv_id,))
        db.execute('DELETE FROM conversations WHERE id=?', (conv_id,))
    return jsonify({'ok': True})

@app.route('/api/chat/history/<int:conv_id>')
@require_auth
def chat_history(conv_id):
    with sqlite3.connect(CHAT_DB) as db:
        rows = db.execute('SELECT role, content FROM messages WHERE conv_id=? ORDER BY id', (conv_id,)).fetchall()
        return jsonify([{'role': r[0], 'content': r[1]} for r in rows])

@app.route('/api/chat', methods=['POST'])
@require_auth
def chat():
    data = request.form.to_dict() if request.form else {}
    user_msg = data.get('message', request.json.get('message', '') if request.is_json else '')
    conv_id = data.get('conv_id', request.json.get('conv_id', 0) if request.is_json else 0)
    conv_id = int(conv_id) if conv_id else 0
    
    # Handle uploaded files in chat
    file_contents = []
    if 'files' in request.files:
        for f in request.files.getlist('files'):
            try:
                if f.filename.lower().endswith('.pdf'):
                    import fitz
                    doc = fitz.open(stream=f.read(), filetype='pdf')
                    text = ''
                    for page in doc: text += page.get_text() + '\n'
                    doc.close()
                    file_contents.append(f'[File: {f.filename}]\n{text[:3000]}')
                elif f.filename.lower().endswith(('.txt', '.md', '.py', '.csv', '.json', '.xml', '.yaml')):
                    content = f.read().decode('utf-8', errors='replace')
                    file_contents.append(f'[File: {f.filename}]\n{content[:5000]}')
                else:
                    file_contents.append(f'[File: {f.filename}] (binary file, not readable)')
            except:
                file_contents.append(f'[File: {f.filename}] (read error)')
    
    if file_contents:
        user_msg = f'[Attached files]\n{chr(10).join(file_contents)}\n\n[User message]\n{user_msg}'
    
    if not user_msg: return jsonify({'reply': 'No message provided.'})
    
    with sqlite3.connect(CHAT_DB) as db:
        db.execute('INSERT INTO messages (conv_id, role, content) VALUES (?,?,?)', (conv_id, 'user', user_msg[:2000]))
        first = db.execute('SELECT COUNT(*) FROM messages WHERE conv_id=?', (conv_id,)).fetchone()[0]
        if first <= 1:
            title = (user_msg[:30] + '...') if len(user_msg) > 30 else user_msg
            db.execute('UPDATE conversations SET title=? WHERE id=?', (title[:50], conv_id))
        db.commit()
        rows = db.execute('SELECT role, content FROM messages WHERE conv_id=? ORDER BY id', (conv_id,)).fetchall()
    
    history = [{'role': r[0], 'content': r[1]} for r in rows]
    
    if DEEPSEEK_KEY:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=DEEPSEEK_KEY, base_url='https://api.deepseek.com/v1')
            resp = client.chat.completions.create(
                model='deepseek-v4-pro',
                messages=[{'role': 'system', 'content': '你是SiC材料实验助手。如果用户上传了文件，仔细阅读文件内容。如果用户要求按文件格式生成内容，严格遵循文件的格式和结构。用中文回答，简洁专业。'}] + history,
                temperature=0.7, max_tokens=800
            )
            reply = resp.choices[0].message.content
        except:
            reply = chat_local(user_msg)
    else:
        reply = chat_local(user_msg)
    
    with sqlite3.connect(CHAT_DB) as db:
        db.execute('INSERT INTO messages (conv_id, role, content) VALUES (?,?,?)', (conv_id, 'assistant', reply))
        db.commit()
    
    return jsonify({'reply': reply})


if __name__ == '__main__':
    print('='*50)
    print('  SiC Lab v3.0 · 密钥保护')
    print(f'  本地: http://localhost:5000')
    print(f'  密钥: {ACCESS_KEY}')
    print('='*50)
    app.run(host='0.0.0.0', port=5000, debug=False)
